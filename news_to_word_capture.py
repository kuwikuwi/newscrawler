import requests
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from PIL import Image
import os
import time
from datetime import datetime
import tempfile

class NewsToWordCapture:
    def __init__(self):
        self.setup_driver()
        self.doc = Document()
        self.setup_document_style()
        
    def setup_driver(self):
        """Chrome 드라이버 설정"""
        chrome_options = Options()
        chrome_options.add_argument('--headless')  # 백그라운드 실행
        chrome_options.add_argument('--window-size=1200,800')  # 적당한 크기
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36')
        
        self.driver = webdriver.Chrome(options=chrome_options)
    
    def setup_document_style(self):
        """워드 문서 기본 스타일 설정"""
        # 제목 추가
        title = self.doc.add_heading('양자통신 관련 기사', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 날짜 추가
        date_para = self.doc.add_paragraph()
        date_para.add_run(f'기간 : {datetime.now().strftime("%Y년 %m월 %d일")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 빈 줄
        self.doc.add_paragraph()
    
    def capture_and_add_to_word(self, url, title, source, category="일반"):
        """뉴스 캡처하고 워드에 추가"""
        try:
            print(f"Processing: {title[:50]}...")
            
            # 페이지 로드
            self.driver.get(url)
            time.sleep(3)  # 페이지 로딩 대기
            
            # 광고나 팝업 제거 시도
            self.remove_popups()
            
            # 임시 스크린샷 파일 생성
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                temp_path = temp_file.name
            
            # 스크린샷 촬영
            self.driver.save_screenshot(temp_path)
            
            # 이미지 최적화 (크기 조정)
            self.optimize_image(temp_path)
            
            # 워드 문서에 추가
            self.add_news_to_document(title, source, category, url, temp_path)
            
            # 임시 파일 삭제
            os.unlink(temp_path)
            
            print(f"✓ 완료: {title[:30]}...")
            return True
            
        except Exception as e:
            print(f"✗ 오류 ({title[:30]}): {str(e)}")
            return False
    
    def remove_popups(self):
        """팝업이나 광고 제거 시도"""
        try:
            # 일반적인 팝업 닫기 버튼들
            popup_selectors = [
                '.popup-close', '.modal-close', '.close-btn', 
                '[aria-label="close"]', '[aria-label="닫기"]',
                '.layer-close', '.btn-close'
            ]
            
            for selector in popup_selectors:
                try:
                    elements = self.driver.find_elements(By.CSS_SELECTOR, selector)
                    for element in elements:
                        if element.is_displayed():
                            element.click()
                            time.sleep(0.5)
                except:
                    continue
                    
        except Exception:
            pass  # 팝업 제거 실패해도 계속 진행
    
    def optimize_image(self, image_path):
        """이미지 최적화 (크기 조정, 품질 개선)"""
        try:
            with Image.open(image_path) as img:
                # 이미지 크기 조정 (너무 크면 워드에서 보기 어려움)
                max_width = 1000
                if img.width > max_width:
                    ratio = max_width / img.width
                    new_height = int(img.height * ratio)
                    img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
                
                # 품질 개선하여 저장
                img.save(image_path, 'PNG', optimize=True, quality=85)
                
        except Exception as e:
            print(f"이미지 최적화 실패: {e}")
    
    def add_news_to_document(self, title, source, category, url, image_path):
        """워드 문서에 뉴스 항목 추가"""
        # 뉴스 번호 (현재 문서의 항목 수 기준)
        news_count = len([p for p in self.doc.paragraphs if p.text.strip().startswith('▲')])
        
        # 카테고리별 이모지
        category_emoji = {
            "양자컴퓨터": "💻",
            "양자통신": "📡", 
            "양자암호": "🔒",
            "양자소자": "⚛️",
            "일반": "📰"
        }
        
        emoji = category_emoji.get(category, "📰")
        
        # 제목 추가 (굵게, 크게)
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(f"{emoji} {title}")
        title_run.bold = True
        title_run.font.size = Inches(0.16)  # 약간 큰 글씨
        
        # 메타 정보 추가
        meta_para = self.doc.add_paragraph()
        meta_text = f"출처: {source} | 카테고리: {category} | 수집일: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        meta_run = meta_para.add_run(meta_text)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)  # 회색
        
        # 스크린샷 이미지 추가
        try:
            # 이미지를 적절한 크기로 삽입 (페이지 너비의 80% 정도)
            self.doc.add_picture(image_path, width=Inches(6))
            
            # 이미지 캡션
            caption_para = self.doc.add_paragraph()
            caption_run = caption_para.add_run(f"▲ {title[:50]}{'...' if len(title) > 50 else ''} (출처: {source})")
            caption_run.italic = True
            caption_run.font.size = Inches(0.12)  # 작은 글씨
            caption_run.font.color.rgb = RGBColor(100, 100, 100)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            # 이미지 삽입 실패 시 텍스트로 대체
            error_para = self.doc.add_paragraph()
            error_run = error_para.add_run(f"[이미지 로드 실패: {str(e)}]")
            error_run.font.color.rgb = RGBColor(255, 0, 0)
        
        # URL 링크 추가
        url_para = self.doc.add_paragraph()
        url_para.add_run("원문 링크: ").bold = True
        url_para.add_run(url)
        
        # 구분선 추가
        self.doc.add_paragraph("─" * 50)
        self.doc.add_paragraph()  # 빈 줄
    
    def save_document(self, filename=None):
        """워드 문서 저장"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"양자통신_뉴스_{timestamp}.docx"
        
        self.doc.save(filename)
        print(f"워드 문서 저장 완료: {filename}")
        return filename
    
    def close(self):
        """리소스 정리"""
        self.driver.quit()

# 기존 크롤링 시스템과 통합하는 함수
def integrate_with_existing_crawler(news_data_list):
    """기존 크롤링 결과와 통합"""
    word_capture = NewsToWordCapture()
    
    try:
        success_count = 0
        total_count = len(news_data_list)
        
        for i, news in enumerate(news_data_list, 1):
            print(f"\n진행률: {i}/{total_count}")
            
            # 필수 필드 확인
            if not all(key in news for key in ['title', 'link', 'source']):
                print(f"필수 필드 누락: {news}")
                continue
            
            # 스크린샷 캡처 및 워드에 추가
            if word_capture.capture_and_add_to_word(
                news['link'], 
                news['title'], 
                news['source'], 
                news.get('category', '일반')
            ):
                success_count += 1
            
            # 서버 부하 방지
            time.sleep(2)
        
        # 워드 문서 저장
        filename = word_capture.save_document()
        
        print(f"\n=== 완료 ===")
        print(f"총 {total_count}개 중 {success_count}개 성공")
        print(f"저장된 파일: {filename}")
        
        return filename
        
    finally:
        word_capture.close()

# 사용 예시
def main():
    # 예시 뉴스 데이터 (실제로는 기존 크롤링 결과를 사용)
    news_list = [
        {
            "title": "키사이트 테크놀로지스, 256큐비트 양자컴퓨터 제어시스템 공급",
            "link": "https://www.etnews.com/20250604000123",
            "source": "전자신문",
            "category": "양자컴퓨터"
        },
        {
            "title": "POSTECH·KAIST·日NIMS, 전자 에너지 제어 양자 소자 세계 최초 개발",
            "link": "https://www.econoscience.co.kr/news/articleView.html?idxno=59876",
            "source": "이코노미사이언스", 
            "category": "양자소자"
        }
        # 실제로는 크롤링 결과에서 가져올 데이터
    ]
    
    # 워드 문서 생성
    integrate_with_existing_crawler(news_list)

if __name__ == "__main__":
    main()