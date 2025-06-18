# 기존 크롤링 코드에 워드 생성 기능을 추가한 통합 버전

from bs4 import BeautifulSoup
from datetime import datetime, timedelta
import requests
import pandas as pd
import re
from urllib.parse import quote
import urllib3
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from PIL import Image
import os
import time
import tempfile

# SSL 경고 비활성화
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

class IntegratedNewsCrawler:
    def __init__(self):
        self.title_text = []
        self.link_text = []
        self.source_text = []
        self.date_text = []
        self.contents_text = []
        
        # 워드 문서 관련
        self.setup_selenium()
        self.doc = Document()
        self.setup_document_style()
        
        # 결과 저장 경로
        self.RESULT_PATH = 'C:/Users/kibae/Desktop/google_news_crawling/result'
        if not os.path.exists(self.RESULT_PATH):
            os.makedirs(self.RESULT_PATH)
    
    def setup_selenium(self):
        """셀레니움 드라이버 설정"""
        chrome_options = Options()
        chrome_options.add_argument('--headless')
        chrome_options.add_argument('--window-size=1200,800')
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36')
        
        self.driver = webdriver.Chrome(options=chrome_options)
    
    def setup_document_style(self):
        """워드 문서 기본 스타일 설정"""
        title = self.doc.add_heading('양자통신 관련 기사', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = self.doc.add_paragraph()
        date_para.add_run(f'기간 : {datetime.now().strftime("%Y년 %m월 %d일")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
    
    def date_cleansing(self, date_str):
        """날짜 정제 함수"""
        now = datetime.now()
        
        try:
            if 'hour' in date_str or 'hr' in date_str:
                hours_ago = int(re.search(r'(\d+)', date_str).group(1))
                target_date = now - timedelta(hours=hours_ago)
                return target_date.strftime('%Y.%m.%d.')
                
            elif 'day' in date_str:
                days_ago = int(re.search(r'(\d+)', date_str).group(1))
                target_date = now - timedelta(days=days_ago)
                return target_date.strftime('%Y.%m.%d.')
                
            elif 'min' in date_str:
                minutes_ago = int(re.search(r'(\d+)', date_str).group(1))
                target_date = now - timedelta(minutes=minutes_ago)
                return target_date.strftime('%Y.%m.%d.')
                
            # 추가 날짜 형식 처리...
            
        except Exception as e:
            print(f"Date conversion error: {e} - Input: {date_str}")
        
        return now.strftime('%Y.%m.%d.')
    
    def contents_cleansing(self, contents):
        """컨텐츠 정제 함수"""
        if contents:
            cleaned_content = re.sub('<.+?>', '', str(contents)).strip()
            cleaned_content = re.sub(r'\s+', ' ', cleaned_content)
            self.contents_text.append(cleaned_content)
        else:
            self.contents_text.append("")
    
    def crawler(self, maxpage, query, time_range=None, language='en'):
        """뉴스 크롤링 (기존 방식)"""
        # 기존 리스트 초기화
        self.title_text.clear()
        self.link_text.clear()
        self.source_text.clear()
        self.date_text.clear()
        self.contents_text.clear()
        
        encoded_query = quote(query)
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept-Language': 'en-US,en;q=0.9',
        }
        
        time_params = {
            'day': 'd',
            'week': 'w', 
            'month': 'm',
            'year': 'y'
        }
        
        time_param = ""
        if time_range in time_params:
            time_param = f"&tbs=qdr:{time_params[time_range]}"
        
        page = 0
        items_per_page = 10
        max_items = int(maxpage) * items_per_page
        
        while page * items_per_page < max_items:
            url = f"https://www.google.com/search?q={encoded_query}&tbm=nws&hl={language}{time_param}&start={page * items_per_page}"
            
            try:
                response = requests.get(url, headers=headers, timeout=30, verify=False)
                if response.status_code != 200:
                    print(f"Failed to fetch page {page+1}. Status code: {response.status_code}")
                    break
                    
                html = response.text
                soup = BeautifulSoup(html, 'html.parser')
                
                # 뉴스 아티클 찾기 (기존 로직)
                articles = soup.select('div.SoaBEf') or soup.select('div.xuvV6b') or soup.select('div.v7W49e')
                
                if not articles:
                    print(f"No articles found on page {page+1}")
                    break
                
                for article in articles:
                    # 제목 추출
                    title_element = (article.select_one('div.MBeuO') or 
                                    article.select_one('h3') or 
                                    article.select_one('.DY5T1d'))
                    
                    if title_element:
                        title = title_element.get_text(strip=True)
                        self.title_text.append(title)
                    else:
                        continue
                    
                    # 링크 추출
                    link_element = article.select_one('a')
                    if link_element and 'href' in link_element.attrs:
                        link = link_element['href']
                        if link.startswith('/url?'):
                            match = re.search(r'url=([^&]+)', link)
                            if match:
                                from urllib.parse import unquote
                                link = unquote(match.group(1))
                        self.link_text.append(link)
                    else:
                        continue
                    
                    # 출처 추출
                    source_element = article.select_one('div.CEMjEf span') or article.select_one('.vN4Yjc')
                    if source_element:
                        source_name = source_element.text.strip()
                        for separator in [' · ', ' - ', ' | ']:
                            if separator in source_name:
                                source_name = source_name.split(separator)[0].strip()
                                break
                        self.source_text.append(source_name)
                    else:
                        self.source_text.append("Unknown")
                    
                    # 날짜 추출
                    date_element = article.select_one('div.OSrXXb span')
                    if date_element:
                        self.date_text.append(self.date_cleansing(date_element.text.strip()))
                    else:
                        self.date_text.append(datetime.now().strftime('%Y.%m.%d.'))
                    
                    # 내용 추출
                    content_element = article.select_one('div.GI74Re')
                    if content_element:
                        self.contents_cleansing(content_element.text)
                    else:
                        self.contents_text.append("")
                
                print(f"Processed page {page+1}")
                page += 1
                time.sleep(2)
                
            except Exception as e:
                print(f"Error crawling page {page+1}: {e}")
                break
    
    def capture_screenshot_and_add_to_word(self, url, title, source, category="일반"):
        """스크린샷 캡처하고 워드에 추가"""
        try:
            print(f"Capturing: {title[:50]}...")
            
            # 페이지 로드
            self.driver.get(url)
            time.sleep(3)
            
            # 임시 스크린샷 파일
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                temp_path = temp_file.name
            
            # 스크린샷 촬영
            self.driver.save_screenshot(temp_path)
            
            # 이미지 최적화
            self.optimize_image(temp_path)
            
            # 워드에 추가
            self.add_to_word_document(title, source, category, url, temp_path)
            
            # 임시 파일 삭제
            os.unlink(temp_path)
            
            return True
            
        except Exception as e:
            print(f"Screenshot error for {title[:30]}: {e}")
            return False
    
    def optimize_image(self, image_path):
        """이미지 최적화"""
        try:
            with Image.open(image_path) as img:
                max_width = 1000
                if img.width > max_width:
                    ratio = max_width / img.width
                    new_height = int(img.height * ratio)
                    img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
                
                img.save(image_path, 'PNG', optimize=True, quality=85)
        except Exception as e:
            print(f"Image optimization failed: {e}")
    
    def add_to_word_document(self, title, source, category, url, image_path):
        """워드 문서에 뉴스 추가"""
        # 카테고리별 이모지
        category_emoji = {
            "양자컴퓨터": "💻",
            "양자통신": "📡",
            "양자암호": "🔒", 
            "양자소자": "⚛️",
            "일반": "📰"
        }
        
        emoji = category_emoji.get(category, "📰")
        
        # 제목
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(f"{emoji} {title}")
        title_run.bold = True
        title_run.font.size = Inches(0.16)
        
        # 메타 정보
        meta_para = self.doc.add_paragraph()
        meta_text = f"출처: {source} | 카테고리: {category} | 수집일: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        meta_run = meta_para.add_run(meta_text)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 스크린샷
        try:
            self.doc.add_picture(image_path, width=Inches(6))
            
            # 캡션
            caption_para = self.doc.add_paragraph()
            caption_run = caption_para.add_run(f"▲ {title[:50]}{'...' if len(title) > 50 else ''} (출처: {source})")
            caption_run.italic = True
            caption_run.font.size = Inches(0.12)
            caption_run.font.color.rgb = RGBColor(100, 100, 100)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            error_para = self.doc.add_paragraph()
            error_run = error_para.add_run(f"[이미지 로드 실패: {str(e)}]")
            error_run.font.color.rgb = RGBColor(255, 0, 0)
        
        # URL
        url_para = self.doc.add_paragraph()
        url_para.add_run("원문 링크: ").bold = True
        url_para.add_run(url)
        
        # 구분선
        self.doc.add_paragraph("─" * 50)
        self.doc.add_paragraph()
    
    def save_all_results(self, query):
        """모든 결과 저장 (엑셀 + 워드)"""
        now = datetime.now()
        timestamp = f'{now.year}-{now.month}-{now.day} {now.hour}시 {now.minute}분 {now.second}초'
        
        # 1. 기존 방식 엑셀 저장
        min_length = min(len(self.title_text), len(self.link_text), 
                        len(self.source_text), len(self.date_text), len(self.contents_text))
        
        result = {
            "contents": self.contents_text[:min_length],
            "title": self.title_text[:min_length],
            "link": self.link_text[:min_length],
            "date": self.date_text[:min_length],
            "source": self.source_text[:min_length],
        }
        
        df = pd.DataFrame(result)
        excel_filename = f'{timestamp} {query}.xlsx'
        excel_path = os.path.join(self.RESULT_PATH, excel_filename)
        df.to_excel(excel_path, sheet_name='sheet1')
        print(f"엑셀 파일 저장: {excel_path}")
        
        # 2. 워드 문서 저장
        word_filename = f'{timestamp} {query}_screenshots.docx'
        word_path = os.path.join(self.RESULT_PATH, word_filename)
        self.doc.save(word_path)
        print(f"워드 파일 저장: {word_path}")
        
        return df, excel_path, word_path
    
    def run_complete_crawling(self, maxpage, query, time_range=None, language='en', 
                            capture_screenshots=True, max_screenshots=10):
        """완전한 크롤링 실행 (텍스트 + 스크린샷)"""
        print("=== 뉴스 크롤링 시작 ===")
        
        # 1단계: 텍스트 크롤링
        print("1단계: 텍스트 데이터 크롤링...")
        self.crawler(maxpage, query, time_range, language)
        
        if not self.title_text:
            print("크롤링된 뉴스가 없습니다.")
            return None
        
        print(f"총 {len(self.title_text)}개 뉴스 발견")
        
        # 2단계: 스크린샷 캡처 (선택사항)
        if capture_screenshots:
            print("2단계: 스크린샷 캡처 중...")
            screenshot_count = 0
            max_capture = min(len(self.title_text), max_screenshots)
            
            for i in range(max_capture):
                if screenshot_count >= max_screenshots:
                    break
                    
                try:
                    # 유효한 URL인지 확인
                    url = self.link_text[i]
                    if not url.startswith('http'):
                        continue
                        
                    if self.capture_screenshot_and_add_to_word(
                        url, 
                        self.title_text[i], 
                        self.source_text[i],
                        "양자통신"  # 기본 카테고리
                    ):
                        screenshot_count += 1
                    
                    time.sleep(2)  # 서버 부하 방지
                    
                except Exception as e:
                    print(f"스크린샷 오류 {i+1}: {e}")
                    continue
            
            print(f"스크린샷 캡처 완료: {screenshot_count}개")
        
        # 3단계: 결과 저장
        print("3단계: 결과 저장...")
        df, excel_path, word_path = self.save_all_results(query)
        
        print("=== 크롤링 완료 ===")
        print(f"엑셀 파일: {excel_path}")
        print(f"워드 파일: {word_path}")
        
        return {
            'dataframe': df,
            'excel_path': excel_path, 
            'word_path': word_path,
            'total_news': len(self.title_text),
            'screenshots_captured': screenshot_count if capture_screenshots else 0
        }
    
    def close(self):
        """리소스 정리"""
        try:
            self.driver.quit()
        except:
            pass

# 사용 예시
def main():
    crawler = IntegratedNewsCrawler()
    
    try:
        # 크롤링 실행
        result = crawler.run_complete_crawling(
            maxpage=3,              # 크롤링할 페이지 수
            query="양자컴퓨터",       # 검색어
            time_range="week",      # 시간 범위 
            language="ko",          # 언어
            capture_screenshots=True, # 스크린샷 캡처 여부
            max_screenshots=5       # 최대 스크린샷 수
        )
        
        if result:
            print(f"\n최종 결과:")
            print(f"- 총 뉴스: {result['total_news']}개")
            print(f"- 스크린샷: {result['screenshots_captured']}개")
            print(f"- 엑셀 파일: {result['excel_path']}")
            print(f"- 워드 파일: {result['word_path']}")
        
    finally:
        crawler.close()

if __name__ == "__main__":
    main()